"""
Document Processor Module
Extracts text from various document formats:
- PDF files using PyMuPDF (fitz)
- Word documents using python-docx
- Excel files using pandas
- CSV files using pandas
- Plain text files
"""

from typing import List
from pathlib import Path
import re


class DocumentProcessor:
    """
    Processes various document formats and extracts text chunks.
    """
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize document processor with chunking parameters.
        
        Args:
            chunk_size: Target size for text chunks (in characters)
            chunk_overlap: Overlap between consecutive chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def process_file(self, file_path: str, filename: str) -> List[str]:
        """
        Process a file and extract text chunks based on file type.
        
        Args:
            file_path: Path to the file
            filename: Original filename
            
        Returns:
            list: List of text chunks
        """
        file_ext = Path(filename).suffix.lower()
        
        try:
            if file_ext == '.pdf':
                text = self._extract_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                text = self._extract_from_word(file_path)
            elif file_ext in ['.xlsx', '.xls']:
                text = self._extract_from_excel(file_path)
            elif file_ext == '.csv':
                text = self._extract_from_csv(file_path)
            elif file_ext == '.txt':
                text = self._extract_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            # Clean and chunk the text
            text = self._clean_text(text)
            chunks = self._chunk_text(text)
            
            return chunks
            
        except Exception as e:
            print(f"Error processing {filename}: {str(e)}")
            raise
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF using PyMuPDF.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            str: Extracted text
        """
        try:
            import fitz  # PyMuPDF
            
            text = []
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text.append(f"\n--- Page {page_num + 1} ---\n")
                text.append(page.get_text())
            
            doc.close()
            return "\n".join(text)
            
        except ImportError:
            return "Error: PyMuPDF (fitz) not installed. Install with: pip install PyMuPDF"
        except Exception as e:
            raise Exception(f"PDF extraction failed: {str(e)}")
    
    def _extract_from_word(self, file_path: str) -> str:
        """
        Extract text from Word document using python-docx.
        
        Args:
            file_path: Path to Word file
            
        Returns:
            str: Extracted text
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            text = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    text.append(" | ".join(row_text))
            
            return "\n".join(text)
            
        except ImportError:
            return "Error: python-docx not installed. Install with: pip install python-docx"
        except Exception as e:
            raise Exception(f"Word extraction failed: {str(e)}")
    
    def _extract_from_excel(self, file_path: str) -> str:
        """
        Extract text from Excel file using pandas.
        
        Args:
            file_path: Path to Excel file
            
        Returns:
            str: Extracted text
        """
        try:
            import pandas as pd
            
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            text = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                text.append(f"\n--- Sheet: {sheet_name} ---\n")
                
                # Convert dataframe to text representation
                text.append(df.to_string(index=False))
            
            return "\n".join(text)
            
        except ImportError:
            return "Error: pandas or openpyxl not installed. Install with: pip install pandas openpyxl"
        except Exception as e:
            raise Exception(f"Excel extraction failed: {str(e)}")
    
    def _extract_from_csv(self, file_path: str) -> str:
        """
        Extract text from CSV file using pandas.
        
        Args:
            file_path: Path to CSV file
            
        Returns:
            str: Extracted text
        """
        try:
            import pandas as pd
            
            df = pd.read_csv(file_path)
            return df.to_string(index=False)
            
        except ImportError:
            return "Error: pandas not installed. Install with: pip install pandas"
        except Exception as e:
            raise Exception(f"CSV extraction failed: {str(e)}")
    
    def _extract_from_txt(self, file_path: str) -> str:
        """
        Extract text from plain text file.
        
        Args:
            file_path: Path to text file
            
        Returns:
            str: Extracted text
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            raise Exception(f"Text extraction failed: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """
        Clean extracted text by removing extra whitespace and formatting.
        
        Args:
            text: Raw extracted text
            
        Returns:
            str: Cleaned text
        """
        # Remove multiple newlines
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # Remove multiple spaces
        text = re.sub(r' +', ' ', text)
        
        # Remove leading/trailing whitespace from each line
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        return text.strip()
    
    def _chunk_text(self, text: str) -> List[str]:
        """
        Split text into overlapping chunks for better context preservation.
        
        Args:
            text: Text to chunk
            
        Returns:
            list: List of text chunks
        """
        if len(text) <= self.chunk_size:
            return [text] if text.strip() else []
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # Try to break at sentence or paragraph boundary
            if end < len(text):
                # Look for paragraph break
                break_point = text.rfind('\n\n', start, end)
                if break_point == -1:
                    # Look for sentence break
                    break_point = text.rfind('. ', start, end)
                if break_point == -1:
                    # Look for any space
                    break_point = text.rfind(' ', start, end)
                if break_point != -1 and break_point > start:
                    end = break_point + 1
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start with overlap
            start = end - self.chunk_overlap
            if start < 0:
                start = 0
        
        return chunks
